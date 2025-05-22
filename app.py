import streamlit as st
import tempfile
import cv2
import whisper
from docx import Document
from docx.shared import Inches
from skimage.metrics import structural_similarity as ssim
import numpy as np
import os

st.title("Video to Word Document Converter")
st.write("Upload a video, and get a Word file with screenshots and voice narration.")

uploaded_file = st.file_uploader("Upload a video", type=["mp4", "mov", "avi", "mkv"])

def extract_screenshots(video_path):
    cap = cv2.VideoCapture(video_path)
    ret, prev = cap.read()
    screenshots = []
    times = []
    
    frame_rate = cap.get(cv2.CAP_PROP_FPS)
    success = True
    frame_number = 0

    while success:
        success, curr = cap.read()
        if not success or curr is None:
            break

        if frame_number % int(frame_rate) == 0:
            grayA = cv2.cvtColor(prev, cv2.COLOR_BGR2GRAY)
            grayB = cv2.cvtColor(curr, cv2.COLOR_BGR2GRAY)
            score, _ = ssim(grayA, grayB, full=True)
            if score < 0.94:
                time_sec = cap.get(cv2.CAP_PROP_POS_MSEC) / 1000.0
                screenshots.append(curr.copy())
                times.append(time_sec)
            prev = curr.copy()
        frame_number += 1

    cap.release()
    return screenshots, times

def generate_doc(screenshots, captions):
    doc = Document()
    for i, (img, text) in enumerate(zip(screenshots, captions)):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_img:
            cv2.imwrite(tmp_img.name, img)
            doc.add_picture(tmp_img.name, width=Inches(5))
            doc.add_paragraph(text)
            os.unlink(tmp_img.name)
    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_doc.name)
    return temp_doc.name

if uploaded_file is not None:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as tmp_vid:
        tmp_vid.write(uploaded_file.read())
        video_path = tmp_vid.name

    st.info("Extracting screenshots...")
    screenshots, times = extract_screenshots(video_path)

    st.success(f"{len(screenshots)} screenshots captured.")
    
    st.info("Transcribing audio...")
    model = whisper.load_model("base")
    result = model.transcribe(video_path)

    # Map transcript segments to screenshot times
    captions = []
    for t in times:
        matched = ""
        for segment in result["segments"]:
            if segment['start'] <= t <= segment['end']:
                matched = segment['text']
                break
        captions.append(matched or "[No speech detected]")

    st.info("Generating Word document...")
    docx_path = generate_doc(screenshots, captions)

    with open(docx_path, "rb") as file:
        st.download_button("Download Word Document", file, file_name="video_notes.docx")